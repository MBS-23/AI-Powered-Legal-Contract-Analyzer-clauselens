"""
docx_parser.py
--------------
Extracts text from Word contracts (.docx).

We use python-docx rather than docx2txt as the primary path because it
lets us keep paragraph-level structure (useful later for detecting
numbered headings like "8. LIABILITY"). docx2txt is kept as a fallback
for files python-docx can't open (rare, but some DOCX exports from legal
software are slightly malformed).

DOCX has no concept of "pages" the way a PDF does (pagination depends on
the viewer/printer), so we don't fabricate page numbers here — clause
extraction will use paragraph/section numbers as the citation unit instead.
"""

from dataclasses import dataclass
from pathlib import Path

import docx2txt
from docx import Document as DocxDocument


@dataclass
class ParsedDocxDocument:
    source_path: str
    paragraphs: list[str]  # non-empty paragraphs, in order

    @property
    def full_text(self) -> str:
        return "\n".join(self.paragraphs)

    @property
    def paragraph_count(self) -> int:
        return len(self.paragraphs)


def extract_text_from_docx(path: str | Path) -> ParsedDocxDocument:
    """Extract paragraph text from a DOCX contract.

    Tries python-docx first (keeps structure); falls back to docx2txt
    (more forgiving parser, but returns one flat text blob) if python-docx
    fails to open the file.
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"No such file: {path}")

    try:
        doc = DocxDocument(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if not paragraphs:
            raise ValueError("python-docx returned no paragraphs")
        return ParsedDocxDocument(source_path=str(path), paragraphs=paragraphs)
    except Exception:
        # Fallback path: docx2txt is more tolerant of odd DOCX files but
        # only gives us one big string, so we split on blank lines to
        # approximate paragraphs.
        raw = docx2txt.process(str(path))
        paragraphs = [line.strip() for line in raw.splitlines() if line.strip()]
        if not paragraphs:
            raise ValueError(f"Could not extract any text from '{path.name}'")
        return ParsedDocxDocument(source_path=str(path), paragraphs=paragraphs)
